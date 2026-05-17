import os
import io
import json
import fitz  # PyMuPDF
import docx  # python-docx

from groq import Groq

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv

# ── Setup ──────────────────────────────────────────

load_dotenv()

groq_client = Groq(
    api_key=os.getenv("GROQ_API_KEY", "")
)

app = FastAPI(title="LexGuard API v2")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Helpers ─────────────────────────────────────────

def ask(prompt: str) -> str:

    res = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.3,
    )

    return res.choices[0].message.content.strip()


def safe_json(text: str):

    text = text.strip()

    if "```" in text:
        text = text.split("```")[1]

        if text.startswith("json"):
            text = text[4:]

    return json.loads(text.strip())


def extract_pdf_text(data: bytes) -> str:

    doc = fitz.open(stream=data, filetype="pdf")

    return "".join(page.get_text() for page in doc)[:14000]


def extract_docx_text(data: bytes) -> str:
    """Extract text from DOCX files."""

    doc = docx.Document(io.BytesIO(data))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Extract table text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)[:14000]


def extract_text(data: bytes, filename: str) -> str:
    """Route extraction by file type."""

    fname = filename.lower()

    if fname.endswith(".pdf"):
        return extract_pdf_text(data)

    elif fname.endswith(".docx"):
        return extract_docx_text(data)

    else:
        # txt / doc fallback
        return data.decode("utf-8", errors="ignore")[:14000]


# ── The 5 Agents ─────────────────────────────────────

def agent_extractor(doc_text: str) -> list:

    prompt = f"""
You are a legal clause extraction expert.

Extract every distinct clause from the document below.

Return ONLY a valid JSON array.

Each item must contain:
- "id"
- "title"
- "text"

No markdown.
No explanation.
Raw JSON only.

DOCUMENT:
{doc_text}
"""

    return safe_json(ask(prompt))


def agent_attacker(clauses: list) -> list:

    prompt = f"""
You are an aggressive consumer-rights lawyer.

For each clause add:
- "attack"

Explain:
- hidden risks
- exploitative wording
- dangers to the signer

Return SAME JSON array.

No markdown.
Raw JSON only.

CLAUSES:
{json.dumps(clauses)}
"""

    return safe_json(ask(prompt))


def agent_defender(clauses: list) -> list:

    prompt = f"""
You are a corporate lawyer defending standard contract clauses.

For each clause add:
- "defense"

Explain why the clause may be:
- standard
- reasonable
- legally necessary

Return SAME JSON array.

No markdown.
Raw JSON only.

CLAUSES:
{json.dumps(clauses)}
"""

    return safe_json(ask(prompt))


def agent_ambiguity(clauses: list) -> list:

    prompt = f"""
You are a legal linguistics expert specializing in ambiguous contract language.

For each clause analyze:
1. vague wording
2. undefined terms
3. loopholes
4. contradictions

Add:

"ambiguity": {{
    "vague_terms": [],
    "has_contradiction": false,
    "contradiction_note": "",
    "ambiguity_score": 0,
    "plain_warning": ""
}}

Return SAME JSON array.

No markdown.
Raw JSON only.

CLAUSES:
{json.dumps(clauses)}
"""

    return safe_json(ask(prompt))


def agent_judge(clauses: list) -> list:

    prompt = f"""
You are a neutral senior judge with 30 years of contract law experience.

For each clause add:
- "risk_level"
- "risk_score"
- "verdict"
- "recommendation"
- "benchmark"
- "scenario"

Rules:

risk_level:
"HIGH", "MEDIUM", or "LOW"

recommendation:
"REJECT", "NEGOTIATE", or "ACCEPT"

benchmark:
Compare against industry standards.

scenario:
Explain one real-world consequence using "you".

Return SAME JSON array.

No markdown.
Raw JSON only.

CLAUSES:
{json.dumps(clauses)}
"""

    return safe_json(ask(prompt))


# ── Core Logic ───────────────────────────────────────

def run_analysis(text: str, filename: str) -> dict:

    clauses = agent_extractor(text)

    attacked = agent_attacker(clauses)

    defended = agent_defender(attacked)

    flagged = agent_ambiguity(defended)

    final = agent_judge(flagged)

    scores = [c.get("risk_score", 50) for c in final]

    overall_score = (
        int(sum(scores) / len(scores))
        if scores else 50
    )

    overall_level = (
        "HIGH"
        if overall_score > 65
        else "MEDIUM"
        if overall_score > 35
        else "LOW"
    )

    ambiguous_count = sum(
        1
        for c in final
        if c.get("ambiguity", {}).get("ambiguity_score", 0) > 40
    )

    return {
        "success": True,
        "document_name": filename,
        "overall_risk_score": overall_score,
        "overall_risk_level": overall_level,
        "total_clauses": len(final),
        "high_risk_count": sum(
            1 for c in final if c.get("risk_level") == "HIGH"
        ),
        "reject_count": sum(
            1 for c in final if c.get("recommendation") == "REJECT"
        ),
        "ambiguous_count": ambiguous_count,
        "clauses": final,
    }


# ── Routes ───────────────────────────────────────────

@app.post("/analyze")
async def analyze_file(file: UploadFile = File(...)):

    fname = file.filename or "document"

    allowed = (".pdf", ".docx", ".txt", ".doc")

    if not any(fname.lower().endswith(ext) for ext in allowed):
        raise HTTPException(
            400,
            f"Unsupported file type. Allowed: {', '.join(allowed)}"
        )

    data = await file.read()

    text = extract_text(data, fname)

    if len(text.strip()) < 50:
        raise HTTPException(
            400,
            "Document too short or extraction failed"
        )

    try:
        return run_analysis(text, fname)

    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/analyze-text")
async def analyze_text_endpoint(body: dict):

    text = body.get("text", "").strip()

    if len(text) < 50:
        raise HTTPException(400, "Text too short")

    try:
        return run_analysis(text[:14000], "Pasted Text")

    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/health")
def health():

    return {
        "status": "ok",
        "service": "LexGuard v2",
        "agents": 5
    }